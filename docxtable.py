import io
from typing import Iterable
import docx
import docx.document as docxdocument
from docx.shared import Cm
from fields import Base64ImageStr, get_binary_image_data
from models import Additional, Contacts, Education, Experience, Personal
from models import WorkBlankModel


class Image:
    def __init__(self, base64str: Base64ImageStr) -> None:
        self.data = base64str


TableValue = str | Image
TableRow = tuple[str, TableValue]
TableRows = tuple[TableRow, ...]


class DocxTable:
    def __init__(self, tables: dict[str, Iterable[TableRows]]) -> None:
        self.tables = tables

    @staticmethod
    def from_blank(blank: WorkBlankModel) -> 'DocxTable':
        return DocxTable({
            'Личные данные': (get_personal_rows(blank.personal),),
            'Контакты': (get_contacts_rows(blank.contacts),),
            'Образование': tuple(
                get_education_rows(i) for i in blank.education
            ),
            'Опыт работы': tuple(
                get_experience_rows(i) for i in blank.experience
            ),
            'Дополнительно': (get_additional_rows(blank.additional),)
        })

    def as_docx_stream(self) -> io.BytesIO:
        document: docxdocument.Document = docx.Document()
        for tables_name, tables_data in self.tables.items():
            document.add_heading(tables_name)
            for table_data in tables_data:
                self._insert_table(document, table_data)
        stream = io.BytesIO()
        document.save(stream)
        stream.seek(0)
        return stream

    def _insert_table(self, document: docxdocument.Document, table_data):
        table = document.add_table(len(table_data), 2)
        for row, row_data in zip(table.rows, table_data):
            cells = row.cells
            cells[0].add_paragraph(row_data[0])
            paragraph = cells[1].add_paragraph()  # type: ignore
            run = paragraph.add_run()
            self._insert_table_value(row_data[1], run)
        document.add_page_break()

    def _insert_table_value(self, value: TableValue, run):
        if isinstance(value, str):
            run.add_text(value)
        elif isinstance(value, Image):
            run.add_picture(
                get_binary_image_data(value.data),
                width=Cm(7)
            )


def get_personal_rows(personal: Personal) -> TableRows:
    return (
        ('Имя на русском', personal.russian_name),
        ('Имя латиницей', personal.english_name),
        ('Семейный статус', personal.family_status.to_human_name()),
        ('День рождения', format(personal.birthday, "%d.%m.%Y")),
        ('Желаемая зарплата', str(personal.salary)),
        ('Желаемая работа', personal.wanted_job.to_human_name())
    )


def get_contacts_rows(contacts: Contacts):
    return (
        ('Страна', contacts.country),
        ('Город', contacts.city),
        ('Улица', contacts.street),
        ('Квартира', contacts.flat),
        ('Индекс', contacts.index),
        ('Телефон', contacts.phone),
        ('E-mail', contacts.email)
    )


def get_education_rows(education: Education) -> TableRows:
    return (
        ('Период обучения', f'{education.start}-{education.end}'),
        ('Название ВУЗа', education.name),
        ('Специализация', education.specialization),
        ('Квалификация', education.qualification),
        ('Фото диплома', Image(education.photo))
    )


def get_experience_rows(experience: Experience) -> TableRows:
    return (
        ('Период работы', f'{experience.start}-{experience.end}'),
        ('Организация', experience.organization),
        ('Позиция', experience.position),
        ('Навыки', experience.skills)
    )


def get_additional_rows(additional: Additional) -> TableRows:
    return (
        ('Водительские права', additional.driving_license),
        (
            'Привлекался ли к уголовной ответственности',
            "Да" if additional.had_criminal_liability else "Нет"
        ),
        ('Дополнительно', additional.additional)
    )
